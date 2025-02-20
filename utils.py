import pandas as pd
from docx import Document
from datetime import datetime
import os
import logging

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
console_handler.setFormatter(formatter)
logger.addHandler(console_handler)

def _generate_filename(extension):
    """Generate a filename with a timestamp."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"scraped_data_{timestamp}.{extension}"

def save_as_excel(data):
    """Save scraped data as an Excel file with dynamic column structure."""
    try:
        filename = _generate_filename('xlsx')
        if not data:
            raise ValueError("No data to export")
        
        if isinstance(data[0], str):
            # YouTube comments: single "Comments" column
            df = pd.DataFrame(data, columns=['Comments'])
        elif isinstance(data[0], dict) and 'title' in data[0] and 'content' in data[0]:
            # News articles: "Title" and "Content" columns
            df = pd.DataFrame(data)[['title', 'content']]
            df.columns = ['Title', 'Content']  # Rename for consistency
        else:
            raise ValueError("Unsupported data format")
        
        df.to_excel(filename, index=False, engine='openpyxl')
        logger.info(f"Excel saved as {filename}")
        return filename
    except Exception as e:
        logger.error(f"Excel generation failed: {str(e)}")
        raise

def save_as_word(data):
    """Save scraped data as a Word file with dynamic column structure."""
    try:
        filename = _generate_filename('docx')
        doc = Document()
        doc.add_heading('Scraped Data', level=1)
        doc.add_paragraph(f"Total items: {len(data)}\nDate: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        if not data:
            raise ValueError("No data to export")
        
        if isinstance(data[0], str):
            # YouTube comments: single "Comments" column
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Comments'
            for comment in data:
                row_cells = table.add_row().cells
                row_cells[0].text = comment
        elif isinstance(data[0], dict) and 'title' in data[0] and 'content' in data[0]:
            # News articles: "Title" and "Content" columns
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Title'
            hdr_cells[1].text = 'Content'
            for item in data:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get('title', 'No Title')
                row_cells[1].text = item.get('content', 'No Content')
        else:
            raise ValueError("Unsupported data format")
        
        doc.save(filename)
        logger.info(f"Word document saved as {filename}")
        return filename
    except Exception as e:
        logger.error(f"Word generation failed: {str(e)}")
        raise